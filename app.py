from flask import Flask, render_template, request, jsonify, send_file
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import io
import re
import datetime
import os

app = Flask(__name__)

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/120.0.0.0 Safari/537.36"
)


def is_natgeo_url(url):
    """只允许 natgeo 相关域名，避免被当作通用爬虫。"""
    parsed = urlparse(url)
    domain = parsed.netloc.lower()
    return "nationalgeographic" in domain or domain.endswith("natgeo.com")


def parse_docx(file_storage):
    """解析用户上传的 docx，提取标题、正文等。"""
    data = file_storage.read()
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return {"title": "", "body": ""}

    title = ""
    body_paras = []
    for i, p in enumerate(paragraphs):
        # 常见模板：第一段是 NATIONAL GEOGRAPHIC 杂志名，第二段是标题
        if i == 0 and ("NATIONAL GEOGRAPHIC" in p.upper() or "国家地理" in p):
            continue
        if i <= 1 and not title and not re.match(r"^\[\d+\]", p):
            title = p
            continue
        body_paras.append(p)

    # 如果正文中已有 [1] 编号，去掉编号避免重复
    cleaned = []
    for p in body_paras:
        cleaned.append(re.sub(r"^\[\d+\]\s*", "", p))

    return {"title": title, "body": "\n\n".join(cleaned)}


def fetch_metadata(url):
    """仅抓取公开 meta 信息，不抓取全文正文。"""
    headers = {"User-Agent": USER_AGENT}
    resp = requests.get(url, headers=headers, timeout=15)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")

    def get_meta(name=None, prop=None, attr="content"):
        if name:
            tag = soup.find("meta", attrs={"name": name})
            if tag:
                return tag.get(attr, "")
        if prop:
            tag = soup.find("meta", attrs={"property": prop})
            if tag:
                return tag.get(attr, "")
        return ""

    title = get_meta(prop="og:title") or soup.title.string.strip() if soup.title else ""
    description = get_meta(prop="og:description") or get_meta(name="description")
    image = get_meta(prop="og:image")
    author = get_meta(name="author") or get_meta(prop="og:author") or get_meta(prop="article:author")
    published = get_meta(prop="article:published_time") or get_meta(name="publish_date")

    # 尝试从 JSON-LD 获取更准确的日期/作者
    for script in soup.find_all("script", type="application/ld+json"):
        text = script.string or ""
        if '"datePublished"' in text and not published:
            m = re.search(r'"datePublished"\s*:\s*"([^"]+)"', text)
            if m:
                published = m.group(1)
        if '"author"' in text and not author:
            m = re.search(r'"author"\s*:\s*\{[^}]*"name"\s*:\s*"([^"]+)"', text)
            if m:
                author = m.group(1)

    return {
        "title": title.strip(),
        "description": description.strip(),
        "image": image.strip(),
        "author": author.strip(),
        "published": published.strip()[:10] if published else "",
        "url": url,
    }


def add_numbered_paragraphs(doc, body_text):
    """把正文按空行分段，并像模板一样加上 [1] [2] 编号。

    排版针对打印后写笔记优化：
    - 三号字 (16 pt)
    - 1.5 倍行距，方便行间批注
    - 段后 16 pt，段落之间留空写笔记
    - 两端对齐，充分利用页宽
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body_text) if p.strip()]
    start = 1
    for para in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(f"[{start}] {para}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
        start += 1
    return doc


def build_docx(article):
    doc = Document()

    # 页面边距：适度缩小左右边距，提高正文空间利用率，同时保留打印安全边距
    sections = doc.sections[0]
    sections.top_margin = Inches(0.9)
    sections.bottom_margin = Inches(0.9)
    sections.left_margin = Inches(0.75)
    sections.right_margin = Inches(0.75)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(article.get("title", "Untitled"))
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # 副标题 / 来源信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_parts = []
    if article.get("author"):
        meta_parts.append(f"By {article['author']}")
    if article.get("published"):
        meta_parts.append(article["published"])
    if article.get("url"):
        meta_parts.append(article["url"])
    meta_text = "  |  ".join(meta_parts)
    meta_run = meta.add_run(meta_text)
    meta_run.font.name = "Times New Roman"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()  # 空行

    # 摘要
    if article.get("description"):
        desc = doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc_run = desc.add_run(article["description"])
        desc_run.font.name = "Times New Roman"
        desc_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        desc_run.font.size = Pt(11)
        desc_run.font.italic = True
        desc_run.font.color.rgb = RGBColor(64, 64, 64)
        doc.add_paragraph()

    # 正文
    if article.get("body"):
        add_numbered_paragraphs(doc, article["body"])

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generated by NatGeo Study Reader")
    footer_run.font.name = "Times New Roman"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/fetch", methods=["POST"])
def fetch():
    data = request.get_json() or {}
    url = data.get("url", "").strip()
    if not url:
        return jsonify({"error": "请输入文章链接"}), 400
    if not url.startswith("http"):
        url = "https://" + url
    if not is_natgeo_url(url):
        return jsonify({"error": "仅支持 National Geographic 相关链接"}), 400
    try:
        meta = fetch_metadata(url)
        return jsonify(meta)
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"网络请求失败：{str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"解析失败：{str(e)}"}), 500


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "请选择文件"}), 400
    file = request.files["file"]
    if not file.filename.endswith(".docx"):
        return jsonify({"error": "仅支持 .docx 文件"}), 400
    try:
        result = parse_docx(file.stream)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": f"解析失败：{str(e)}"}), 500


@app.route("/download", methods=["POST"])
def download():
    data = request.get_json() or {}
    if not data.get("title"):
        return jsonify({"error": "缺少标题"}), 400
    buffer = build_docx(data)
    filename = re.sub(r"[^\w\-]+", "_", data["title"])[:50] + ".docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_DEBUG", "false").lower() == "true"
    app.run(host="0.0.0.0", port=port, debug=debug)
